"""
Скрипт для (пере)генерации .docx-шаблонов документов.
Запускается один раз вручную: python documents/build_templates.py
Шаблоны — обычные Word-файлы с полями docxtpl вида {{ field }} и циклами
{%tr for item in items %} ... {%tr endfor %} для строк таблицы.
Если понадобится изменить оформление документов — правьте либо этот
скрипт и перегенерируйте шаблоны, либо открывайте .docx из doc_templates/
напрямую в Word и правьте текст вокруг {{ }} руками (сами плейсхолдеры
трогать нельзя, иначе docxtpl их не найдёт).
"""
import os

from docx import Document as WordDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "doc_templates")


def _set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def build_act_template(filename, title, show_transfer_block=True, show_reason=False, reason_label="Основание"):
    """Универсальный шаблон акта: приёма-передачи / перемещения / возврата / списания."""
    doc = WordDocument()
    _set_base_style(doc)

    _add_title(doc, title)
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.add_run("№ {{ number }}").bold = True
    meta.add_run("   от {{ date }}")

    if show_transfer_block:
        doc.add_paragraph("Передал: {{ from_employee }}          Откуда: {{ from_location }}")
        doc.add_paragraph("Принял: {{ to_employee }}          Куда: {{ to_location }}")

    if show_reason:
        doc.add_paragraph(f"{reason_label}: {{{{ reason }}}}")

    doc.add_paragraph("Комментарий: {{ comment }}")
    doc.add_paragraph()

    doc.add_paragraph("Перечень оборудования:")
    # Структура таблицы для docxtpl: заголовок / отдельная строка-маркер "{%tr for%}"
    # (без данных) / строка с данными {{ }} / отдельная строка-маркер "{%tr endfor%}".
    # Маркерные строки НЕ должны содержать ничего кроме тега — иначе docxtpl
    # удаляет вместе со строкой и сам тег, и соседние данные.
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"

    widths = [Cm(1.2), Cm(3.2), Cm(5.5), Cm(3.5), Cm(2.0)]
    headers = ["№", "Инв. номер", "Наименование", "Серийный номер", "Кол-во"]
    for i, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        _set_cell(cell, text, bold=True)
        cell.width = widths[i]

    for_cells = table.rows[1].cells
    _set_cell(for_cells[0], "{%tr for item in items %}")
    for i, cell in enumerate(for_cells):
        cell.width = widths[i]

    data_cells = table.rows[2].cells
    _set_cell(data_cells[0], "{{ item.index }}")
    _set_cell(data_cells[1], "{{ item.inventory_number }}")
    _set_cell(data_cells[2], "{{ item.name }}")
    _set_cell(data_cells[3], "{{ item.serial_number }}")
    _set_cell(data_cells[4], "{{ item.quantity }}")
    for i, cell in enumerate(data_cells):
        cell.width = widths[i]

    end_cells = table.rows[3].cells
    _set_cell(end_cells[0], "{%tr endfor %}")
    for i, cell in enumerate(end_cells):
        cell.width = widths[i]

    doc.add_paragraph()
    doc.add_paragraph()
    if show_transfer_block:
        doc.add_paragraph("Передал: _________________________ / {{ from_employee }} /")
        doc.add_paragraph()
        doc.add_paragraph("Принял: _________________________ / {{ to_employee }} /")
    else:
        doc.add_paragraph("Составил: _________________________ / {{ to_employee }}{{ from_employee }} /")

    doc.save(os.path.join(TEMPLATE_DIR, filename))


def build_statement_template(filename):
    """Ведомость / список оборудования (используется и для инвентаризационной ведомости,
    и для списков по сотруднику/кабинету/подразделению — заголовок и фильтр задаются на лету)."""
    doc = WordDocument()
    _set_base_style(doc)

    _add_title(doc, "{{ title }}")
    doc.add_paragraph()
    doc.add_paragraph("Дата формирования: {{ date }}")
    doc.add_paragraph("Всего позиций: {{ total_count }}")
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=7)
    table.style = "Table Grid"
    headers = ["№", "Инв. номер", "Наименование", "Категория", "Статус", "Местонахождение", "Ответственный"]
    for cell, text in zip(table.rows[0].cells, headers):
        _set_cell(cell, text, bold=True)

    for_cells = table.rows[1].cells
    _set_cell(for_cells[0], "{%tr for item in items %}")

    data_cells = table.rows[2].cells
    _set_cell(data_cells[0], "{{ item.index }}")
    _set_cell(data_cells[1], "{{ item.inventory_number }}")
    _set_cell(data_cells[2], "{{ item.name }}")
    _set_cell(data_cells[3], "{{ item.category }}")
    _set_cell(data_cells[4], "{{ item.status }}")
    _set_cell(data_cells[5], "{{ item.location }}")
    _set_cell(data_cells[6], "{{ item.employee }}")

    end_cells = table.rows[3].cells
    _set_cell(end_cells[0], "{%tr endfor %}")

    doc.add_paragraph()
    doc.add_paragraph("Составил: _________________________")

    doc.save(os.path.join(TEMPLATE_DIR, filename))


def build_equipment_card_template(filename):
    """Карточка единицы имущества — все поля + краткая история эксплуатации."""
    doc = WordDocument()
    _set_base_style(doc)

    _add_title(doc, "Карточка единицы имущества")
    doc.add_paragraph()

    fields = [
        ("Инвентарный номер", "{{ inventory_number }}"),
        ("Серийный номер", "{{ serial_number }}"),
        ("Наименование", "{{ name }}"),
        ("Модель", "{{ model }}"),
        ("Категория", "{{ category }}"),
        ("Технические характеристики", "{{ specifications }}"),
        ("Дата поступления", "{{ purchase_date }}"),
        ("Балансовая стоимость", "{{ purchase_cost }}"),
        ("Статус", "{{ status }}"),
        ("Текущее местонахождение", "{{ location }}"),
        ("Ответственное лицо", "{{ employee }}"),
        ("Примечания", "{{ notes }}"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(fields):
        _set_cell(table.rows[i].cells[0], label, bold=True)
        _set_cell(table.rows[i].cells[1], value)

    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.add_run("История эксплуатации:").bold = True

    # 4 строки: заголовок, отдельная строка-маркер {%tr for%}, строка данных, строка-маркер {%tr endfor%}
    hist_table = doc.add_table(rows=4, cols=4)
    hist_table.style = "Table Grid"
    headers = ["Дата", "Событие", "Кому/куда", "Комментарий"]
    for cell, text in zip(hist_table.rows[0].cells, headers):
        _set_cell(cell, text, bold=True)

    for_cells = hist_table.rows[1].cells
    _set_cell(for_cells[0], "{%tr for event in history %}")

    data_cells = hist_table.rows[2].cells
    _set_cell(data_cells[0], "{{ event.date }}")
    _set_cell(data_cells[1], "{{ event.type }}")
    _set_cell(data_cells[2], "{{ event.target }}")
    _set_cell(data_cells[3], "{{ event.comment }}")

    end_cells = hist_table.rows[3].cells
    _set_cell(end_cells[0], "{%tr endfor %}")

    doc.save(os.path.join(TEMPLATE_DIR, filename))


def main():
    os.makedirs(TEMPLATE_DIR, exist_ok=True)

    build_act_template(
        "transfer_act.docx", "АКТ приёма-передачи оборудования",
        show_transfer_block=True,
    )
    build_act_template(
        "movement_act.docx", "АКТ внутреннего перемещения оборудования",
        show_transfer_block=True,
    )
    build_act_template(
        "return_act.docx", "АКТ возврата оборудования",
        show_transfer_block=True, show_reason=True, reason_label="Состояние при возврате",
    )
    build_act_template(
        "write_off_act.docx", "АКТ списания оборудования",
        show_transfer_block=False, show_reason=True, reason_label="Причина списания",
    )
    build_statement_template("statement.docx")
    build_equipment_card_template("equipment_card.docx")

    print("Шаблоны сгенерированы в", TEMPLATE_DIR)


if __name__ == "__main__":
    main()
